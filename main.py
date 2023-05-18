from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def discard_images(doc):
    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in paragraph.runs:
                if run.text:
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'
                else:
                    try:
                        paragraph.runs.remove(run)
                    except ValueError:
                        pass
        else:
            paragraphs_to_remove.append(paragraph)

    for paragraph in paragraphs_to_remove:
        try:
            doc.paragraphs.remove(paragraph)
        except ValueError:
            pass


def convert_docx_to_txt(input_file, output_file):
    doc = Document(input_file)

    # Discard images
    discard_images(doc)

    with open(output_file, 'w', encoding='utf-8') as txt_file:
        for paragraph in doc.paragraphs:
            if paragraph.text:
                txt_file.write(paragraph.text + '\n')

    print('Conversion complete. Output file:', output_file)


# Usage example
input_file = r'C:\Users\olegs\OneDrive\Рабочий стол\test.docx'
output_file = r'C:\Users\olegs\OneDrive\Рабочий стол\test.txt'
convert_docx_to_txt(input_file, output_file)
